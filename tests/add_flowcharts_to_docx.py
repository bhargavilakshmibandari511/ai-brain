import docx
from docx.shared import Inches
import os

def add_flowcharts_to_docx(input_path, output_path, images_dir):
    doc = docx.Document(input_path)
    
    # Map headings to image filenames (fcX) and captions
    insert_map = {
        "1. Introduction": ("fc7_user_decision_tree.png", "Figure A1: User Task Routing — Feature Decision Tree"),
        "3. System Architecture and Methodology": ("fc1_system_architecture.png", "Figure A2: Enhanced System Architecture Diagram"),
        "3.2 VLM-Based OCR Reader": ("fc2_vlm_ocr_pipeline.png", "Figure A3: VLM-Based OCR Pipeline"),
        "3.3 Knowledge Base and Document Management": ("fc4_rag_ingestion.png", "Figure A4: RAG Document Ingestion Pipeline"),
        "4. Multi-Agent Orchestration": ("fc3_multi_agent_rag.png", "Figure A5: Multi-Agent Routing & RAG Query Flow"),
        "5. Experimental Evaluation and Results": ("fc6_performance_chart.png", "Figure A6: Performance Benchmark Charts"),
        "6. Discussion": ("fc5_mode_decision.png", "Figure A7: Operational Mode Selection")
    }

    # Iterate through paragraphs once and find insertion points
    # We'll insert AFTER the paragraph containing the heading or the section after it
    for heading_text, (img_filename, caption) in insert_map.items():
        image_path = os.path.join(images_dir, img_filename)
        if not os.path.exists(image_path):
            print(f"Warning: Image {image_path} not found.")
            continue
            
        found = False
        for i, paragraph in enumerate(doc.paragraphs):
            if heading_text in paragraph.text:
                # Insert after the next paragraph to give some context before the image
                # Actually, inserting right after the heading is fine too.
                # Let's insert after paragraph i+1 to be safe and avoid breaking the heading itself
                insert_idx = min(i + 2, len(doc.paragraphs))
                
                # We can't really insert into the list doc.paragraphs easily, we have to use add_picture
                # But python-docx documentation says paragraphs are objects. 
                # To insert at a specific position is tricky.
                # Common way: insert_paragraph_before on the next paragraph.
                
                if insert_idx < len(doc.paragraphs):
                    target_p = doc.paragraphs[insert_idx]
                    # New p for image
                    img_p = target_p.insert_paragraph_before()
                    run = img_p.add_run()
                    run.add_picture(image_path, width=Inches(6.0))
                    
                    # New p for caption
                    caption_p = target_p.insert_paragraph_before()
                    caption_run = caption_p.add_run(caption)
                    caption_run.bold = True
                    caption_run.italic = True
                    caption_p.alignment = docx.enum.text.WD_ALIGN_PARAGRAPH.CENTER
                    
                    found = True
                    print(f"Inserted {img_filename} after '{heading_text}'")
                    break
        
        if not found:
            print(f"Warning: Heading '{heading_text}' not found in document.")

    doc.save(output_path)
    print(f"New document saved to {output_path}")

if __name__ == "__main__":
    input_file = "Offline_AI_Digital_Brain_Research_Paper.docx"
    output_file = "Offline_AI_Digital_Brain_Research_Paper_With_Flowcharts.docx"
    images_folder = "flowchart_images"
    
    add_flowcharts_to_docx(input_file, output_file, images_folder)
